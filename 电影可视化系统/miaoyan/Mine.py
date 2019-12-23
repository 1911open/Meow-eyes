# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'Mine.ui'
#
# Created by: PyQt5 UI code generator 5.11.3
#
# WARNING! All changes made in this file will be lost!

from PyQt5 import QtCore, QtGui, QtWidgets
import Mine1
import Mine2
import Mine3
import Mine4
import Mine5
import os
from docx import Document
from docx.shared import Inches
from pyecharts_snapshot.main import make_a_snapshot

class Ui_Mine(object):

    def click6(self):
        '''
        files = os.listdir(".")  # 列出当前目录下所有的文件
        for filename in files:
            portion = os.path.splitext(filename)  # 分离文件名字和后缀
            print(portion)
            if portion[1] == ".html":  # 根据后缀来修改,如无后缀则空
                newname = portion[0] + ".png"  # 要改的新后缀
                os.rename(filename, newname)
        '''
        doc = Document()
        string = "票房月份占比报表"
        doc.add_paragraph(string)
        for year in range(2015,2019,1):
            for month in range(1,13,1):
                try:
                    exec(doc.add_picture('db%d.png'%(year*100+month), width=Inches(6)))
                except:
                    pass
        string = "票房季度占比报表"
        doc.add_paragraph(string)
        for year in range(2015, 2019, 1):
            try:
                exec(doc.add_picture('db%d.png' % (year * 100 + 123), width=Inches(6)))
            except:
                pass
            try:
                exec(doc.add_picture('db%d.png' % (year * 100 + 456), width=Inches(6)))
            except:
                pass
            try:
                exec(doc.add_picture('db%d.png' % (year * 100 + 789), width=Inches(6)))
            except:
                pass
            try:
                exec(doc.add_picture('db%d.png' % (year * 100 + 101112), width=Inches(6)))
            except:
                pass
        string = "票房榜单报表"
        doc.add_paragraph(string)
        for year in range(2015, 2019, 1):
            try:
                exec(doc.add_picture("wc%d.png" % (year * 100 + 3), width=Inches(6)))
            except:
                pass
            try:
                exec(doc.add_picture("wc%d.png" % (year * 100 + 5), width=Inches(6)))
            except:
                pass
            try:
                exec(doc.add_picture("wc%d.png" % (year * 100 + 10), width=Inches(6)))
            except:
                pass
        string = "票房变化报表"
        doc.add_paragraph(string)
        try:
            exec(doc.add_picture("2015_to_2018.png", width=Inches(6)))
        except:
            pass
        string = "演员劳模报表"
        doc.add_paragraph(string)
        for year in range(2015, 2019, 1):
            try:
                exec(doc.add_picture("t4%d.png" % (year * 100 + 3), width=Inches(6)))
            except:
                pass
            try:
                exec(doc.add_picture("wct%d.png" % (year * 100 + 3), width=Inches(6)))
            except:
                pass
            try:
                exec(doc.add_picture("t4%d.png" % (year * 100 + 5), width=Inches(6)))
            except:
                pass
            try:
                exec(doc.add_picture("wct%d.png" % (year * 100 + 5), width=Inches(6)))
            except:
                pass
            try:
                exec(doc.add_picture("t4%d.png" % (year * 100 + 10), width=Inches(6)))
            except:
                pass
            try:
                exec(doc.add_picture("wct%d.png" % (year * 100 + 10), width=Inches(6)))
            except:
                pass
        doc.save('report.docx')
        os.popen('report.docx')
    def click5(self):
        self.form.close()
        Form1 = QtWidgets.QMainWindow()
        self.ui = Mine5.Ui_Mine5()
        self.ui.setupUi(Form1)
        Form1.show()
    def click4(self):
        self.form.close()
        Form1 = QtWidgets.QMainWindow()
        self.ui = Mine4.Ui_Mine4()
        self.ui.setupUi(Form1)
        Form1.show()
    def click3(self):
        self.form.close()
        Form1 = QtWidgets.QMainWindow()
        self.ui = Mine3.Ui_Main3()
        self.ui.setupUi(Form1)
        Form1.show()
    def click2(self):
        self.form.close()
        Form1 = QtWidgets.QMainWindow()
        self.ui = Mine2.Ui_Mine2()
        self.ui.setupUi(Form1)
        Form1.show()
    def click1(self):
        self.form.close()
        Form1 = QtWidgets.QMainWindow()
        self.ui = Mine1.Ui_Mine1()
        self.ui.setupUi(Form1)
        Form1.show()
    def setupUi(self, Mine):
        Mine.setObjectName("Mine")
        Mine.resize(754, 776)
        Mine.setStyleSheet("border-image: url(:/background1.jpg);")
        self.form = Mine
        self.centralwidget = QtWidgets.QWidget(Mine)
        self.centralwidget.setObjectName("centralwidget")
        self.pushButton = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton.setGeometry(QtCore.QRect(40, 580, 271, 41))
        self.pushButton.setStyleSheet("border-image: url(:/M1.jpg);\n"
"font: 20pt \"幼圆\";\n"
"color: rgb(255, 0, 0);")
        self.pushButton.setObjectName("pushButton")
        self.pushButton_2 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_2.setGeometry(QtCore.QRect(440, 470, 271, 41))
        self.pushButton_2.setStyleSheet("border-image: url(:/M2.jpg);\n"
"color: rgb(255, 0, 0);\n"
"font: 20pt \"幼圆\";")
        self.pushButton_2.setObjectName("pushButton_2")
        self.pushButton_3 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_3.setGeometry(QtCore.QRect(40, 360, 271, 41))
        self.pushButton_3.setStyleSheet("border-image: url(:/M3.jpg);\n"
"color: rgb(255, 0, 0);\n"
"font: 20pt \"幼圆\";")
        self.pushButton_3.setObjectName("pushButton_3")
        self.pushButton_4 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_4.setGeometry(QtCore.QRect(440, 360, 271, 41))
        self.pushButton_4.setStyleSheet("border-image: url(:/M4.jpg);\n"
"color: rgb(255, 0, 0);\n"
"font: 20pt \"幼圆\";")
        self.pushButton_4.setObjectName("pushButton_4")
        self.pushButton_5 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_5.setGeometry(QtCore.QRect(40, 470, 271, 41))
        self.pushButton_5.setStyleSheet("border-image: url(:/M5.jpg);\n"
"color: rgb(255, 0, 0);\n"
"font: 20pt \"幼圆\";")
        self.pushButton_5.setObjectName("pushButton_5")
        self.pushButton_6 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_6.setGeometry(QtCore.QRect(440, 580, 271, 41))
        self.pushButton_6.setStyleSheet("border-image: url(:/M6.jpg);\n"
"font: 20pt \"幼圆\";\n"
"color: rgb(255, 0, 0);")
        self.pushButton_6.setObjectName("pushButton_6")
        self.pushButton_7 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_7.setGeometry(QtCore.QRect(110, 50, 530, 261))
        self.pushButton_7.setStyleSheet("font: 20pt \"幼圆\";\n"
                                        "color: rgb(255, 0, 0);\n"
                                        "border-image: url(:/M16.jpg);")
        self.pushButton_7.setObjectName("pushButton_7")
        Mine.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(Mine)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 754, 26))
        self.menubar.setObjectName("menubar")
        Mine.setMenuBar(self.menubar)

        self.retranslateUi(Mine)
        self.pushButton_6.clicked.connect(Mine.close)
        self.pushButton_3.clicked.connect(self.click1)
        self.pushButton_4.clicked.connect(self.click2)
        self.pushButton_5.clicked.connect(self.click3)
        self.pushButton_2.clicked.connect(self.click4)
        self.pushButton_7.clicked.connect(self.click6)
        self.pushButton.clicked.connect(self.click5)
        QtCore.QMetaObject.connectSlotsByName(Mine)

    def retranslateUi(self, Mine):
        _translate = QtCore.QCoreApplication.translate
        Mine.setWindowTitle(_translate("Mine", "欢迎使用喵眼！"))
        self.pushButton.setText(_translate("Mine", "     数据概览"))
        self.pushButton_2.setText(_translate("Mine", "    演员劳模"))
        self.pushButton_3.setText(_translate("Mine", "    票房占比"))
        self.pushButton_4.setText(_translate("Mine", "    票房榜单"))
        self.pushButton_5.setText(_translate("Mine", "    票房变化"))
        self.pushButton_6.setText(_translate("Mine", "    退出程序"))
        self.pushButton_7.setText(_translate("Mine", "综 合 报 表"))

import myqrc_rc
if __name__=="__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    Mine = QtWidgets.QMainWindow()  # 注意，这里和我们一开始创建窗体时使用的界面类型相同
    ui = Ui_Mine()
    ui.setupUi(Mine)
    Mine.show()
    sys.exit(app.exec_())
